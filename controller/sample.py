from fastapi import HTTPException, Depends, BackgroundTasks, Response
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from models.models import Transcription, Meeting
from schemas.transcriptionSchema import TranscriptionCreate
from database import get_db, SessionLocal
import whisper
import os
from datetime import datetime
from dateutil.parser import parse
import json
from dotenv import load_dotenv
from typing import Dict
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
import tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import google.generativeai as genai
import requests

load_dotenv()

# Configure Gemini API
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Load Whisper Model (for Speech-to-Text)
whisper_model = whisper.load_model("tiny")

# Dictionary to store processing status
processing_status: Dict[str, dict] = {}


def translate_meeting_content(meeting_data: dict, target_language: str) -> dict:
    """
    Translate meeting content dynamically using Gemini API
    """
    if target_language.lower() not in ["marathi", "english"]:
        return meeting_data
    
    # If already in target language, return as-is
    current_language = meeting_data.get("language", "english")
    if current_language.lower() == target_language.lower():
        return meeting_data
    
    try:
        # Prepare content for translation - INCLUDING TRANSCRIPT
        content_to_translate = {
            "title": meeting_data.get("title", ""),
            "summary": meeting_data.get("summary", ""),
            "key_points": meeting_data.get("key_points", []),
            "action_items": meeting_data.get("action_items", []),
            "participants": meeting_data.get("participants", []),
            "transcript": meeting_data.get("transcript", "")
        }
        
        if target_language.lower() == "marathi":
            prompt = f"""
Translate the following meeting content to Marathi. Maintain the JSON structure exactly.
Return only valid JSON without any markdown formatting or code blocks.

Content to translate:
{json.dumps(content_to_translate, ensure_ascii=False)}

Requirements:
- Translate all text fields to Marathi including the full transcript
- Keep the same JSON structure
- Translate title, summary, transcript, all items in key_points, action_items arrays
- For participants, translate only if they are generic terms; keep proper names unchanged
- Return pure JSON only
"""
        else:
            prompt = f"""
Translate the following meeting content to English. Maintain the JSON structure exactly.
Return only valid JSON without any markdown formatting or code blocks.

Content to translate:
{json.dumps(content_to_translate, ensure_ascii=False)}

Requirements:
- Translate all text fields to English including the full transcript
- Keep the same JSON structure
- Translate title, summary, transcript, all items in key_points, action_items arrays
- For participants, translate only if they are generic terms; keep proper names unchanged
- Return pure JSON only
"""
        
        # Get available models
        available_models = []
        try:
            for model in genai.list_models():
                if 'generateContent' in model.supported_generation_methods:
                    available_models.append(model.name)
        except:
            pass
        
        # Choose model
        preferred_models = [
            'models/gemini-2.5-flash',
            'models/gemini-2.0-flash',
            'models/gemini-flash-latest'
        ]
        
        model_name = None
        for preferred in preferred_models:
            if preferred in available_models:
                model_name = preferred
                break
        
        if not model_name and available_models:
            model_name = available_models[0]
        
        if not model_name:
            raise Exception("No models available")
        
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        content = response.text
        
        # Clean the response
        if content.startswith("```json"):
            content = content.replace("```json", "").replace("```", "").strip()
        elif content.startswith("```"):
            content = content.replace("```", "").strip()
        
        translated_data = json.loads(content)
        
        # Update meeting data with translations - INCLUDING TRANSCRIPT
        meeting_data["title"] = translated_data.get("title", meeting_data["title"])
        meeting_data["summary"] = translated_data.get("summary", meeting_data["summary"])
        meeting_data["key_points"] = translated_data.get("key_points", meeting_data["key_points"])
        meeting_data["action_items"] = translated_data.get("action_items", meeting_data["action_items"])
        meeting_data["participants"] = translated_data.get("participants", meeting_data["participants"])
        meeting_data["transcript"] = translated_data.get("transcript", meeting_data["transcript"])
        meeting_data["language"] = target_language.lower()
        
        return meeting_data
        
    except Exception as e:
        print(f"Translation error: {e}")
        # Return original data if translation fails
        return meeting_data


async def process_upload_background(file_content: bytes, filename: str, task_id: str):
    db = SessionLocal()
    try:
        processing_status[task_id] = {
            "status": "processing",
            "progress": 0,
            "filename": filename
        }

        now = datetime.now()
        date_clean = now.strftime("%m-%d-%YT%H-%M-%S")
        directory = f"recordings/{date_clean}"
        os.makedirs(directory, exist_ok=True)
        file_path = f"{directory}/{filename}"

        with open(file_path, "wb") as f:
            f.write(file_content)

        processing_status[task_id]["progress"] = 20

        result = whisper_model.transcribe(file_path, fp16=False)
        processing_status[task_id]["progress"] = 50

        transcription_data = TranscriptionCreate(
            transcript=result["text"],
            file_name=filename,
        )
        stored_transcription = create_transcription(transcription_data, db)
        processing_status[task_id]["progress"] = 70

        transcript = process_transcirption(stored_transcription.id, db)
        processing_status[task_id]["progress"] = 90

        meeting_id = transcript.get("id")
        processing_status[task_id] = {
            "status": "completed",
            "progress": 100,
            "result": transcript,
            "meeting_id": meeting_id,
            "filename": filename
        }

    except Exception as e:
        processing_status[task_id] = {
            "status": "error",
            "error": str(e),
            "filename": filename
        }
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()


async def process_upload(file, session: Session, background_tasks: BackgroundTasks):
    try:
        task_id = f"task_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        processing_status[task_id] = {
            "status": "pending",
            "progress": 0,
            "filename": file.filename
        }

        file_content = await file.read()

        background_tasks.add_task(
            process_upload_background,
            file_content,
            file.filename,
            task_id
        )

        return {
            "task_id": task_id,
            "message": "Processing started",
            "status": "pending"
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def get_processing_status(task_id: str):
    if task_id not in processing_status:
        raise HTTPException(status_code=404, detail="Task not found")
    return processing_status[task_id]


def create_transcription(transcription: TranscriptionCreate, db: Session):
    db_transcription = Transcription(**transcription.dict())
    db.add(db_transcription)
    db.commit()
    db.refresh(db_transcription)
    return db_transcription


def call_gemini_api_directly(prompt: str, model_name: str = "models/gemini-2.5-flash") -> str:
    api_key = os.getenv("GOOGLE_API_KEY")
    model_for_url = model_name.replace('models/', '')
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_for_url}:generateContent?key={api_key}"
    
    headers = {
        'Content-Type': 'application/json',
    }
    
    data = {
        "contents": [
            {
                "parts": [
                    {
                        "text": prompt
                    }
                ]
            }
        ]
    }
    
    response = requests.post(url, headers=headers, json=data)
    
    if response.status_code == 200:
        result = response.json()
        if 'candidates' in result and len(result['candidates']) > 0:
            return result['candidates'][0]['content']['parts'][0]['text']
        else:
            raise Exception("No content in response")
    else:
        raise Exception(f"API call failed: {response.status_code} - {response.text}")


def process_transcirption(transcription_id: int, db: Session, target_language: str = "english"):
    transcription = db.query(Transcription).filter(Transcription.id == transcription_id).first()
    if not transcription:
        raise HTTPException(status_code=404, detail="Transcription not found")

    if target_language.lower() == "marathi":
        prompt = f"""
मीटिंगचा ट्रान्सक्रिप्ट वाचा आणि सारांश काढा. फक्त JSON आउटपुट द्या.

ट्रान्सक्रिप्ट:
{transcription.transcript}

मराठीत प्रदान करा:
- title (शीर्षक)
- key_points (मुख्य मुद्दे - array)
- action_items (कृती आयटम - array)
- summary (सारांश)
- participants (सहभागी - array)

Return only valid JSON without markdown.
"""
    else:
        prompt = f"""
Read the meeting transcript and extract a summary. Output JSON only.

Transcript:
{transcription.transcript}

Provide:
- title
- key_points (bullet points array)
- action_items (array)
- summary
- participants (array)

Return only valid JSON without markdown.
"""

    try:
        available_models = []
        try:
            for model in genai.list_models():
                if 'generateContent' in model.supported_generation_methods:
                    available_models.append(model.name)
        except Exception as list_error:
            print(f"Error listing models: {list_error}")
        
        if not available_models:
            raise HTTPException(
                status_code=500, 
                detail="No models available. Please check your Google API key and permissions."
            )
        
        preferred_models = [
            'models/gemini-2.5-flash',
            'models/gemini-2.0-flash',
            'models/gemini-flash-latest',
            'models/gemini-2.5-pro'
        ]
        
        model_name = None
        for preferred in preferred_models:
            if preferred in available_models:
                model_name = preferred
                break
        
        if not model_name:
            model_name = available_models[0]
        
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        content = response.text

    except Exception as e:
        try:
            content = call_gemini_api_directly(prompt, model_name)
        except Exception as direct_error:
            error_detail = f"Both SDK and direct API failed. SDK error: {str(e)}. Direct API error: {str(direct_error)}"
            raise HTTPException(status_code=500, detail=error_detail)

    try:
        if content.startswith("```json"):
            content = content.replace("```json", "").replace("```", "").strip()
        elif content.startswith("```"):
            content = content.replace("```", "").strip()
        
        structured_data = json.loads(content)
        title = structured_data.get("title", "Untitled Meeting")
        key_points = structured_data.get("key_points", [])
        action_items = structured_data.get("action_items", [])
        summary = structured_data.get("summary", "No summary")
        participants = structured_data.get("participants", [])

    except json.JSONDecodeError as e:
        title = "Meeting Summary"
        key_points = ["Failed to extract structured data"]
        action_items = ["Review transcription manually"]
        summary = content[:500] if content else "Failed to generate summary"
        participants = []

    key_points_str = json.dumps(key_points, ensure_ascii=False)
    action_items_str = json.dumps(action_items, ensure_ascii=False)
    participants_str = json.dumps(participants, ensure_ascii=False)

    new_meeting = Meeting(
        title=title,
        date=datetime.now(),
        transcript_id=transcription_id,
        key_points=key_points_str,
        action_items=action_items_str,
        summary=summary,
        participants=participants_str,
        created_at=datetime.now()
    )

    db.add(new_meeting)
    db.commit()
    db.refresh(new_meeting)

    return {
        "id": new_meeting.id,
        "title": new_meeting.title,
        "date": new_meeting.date,
        "transcript_id": new_meeting.transcript_id,
        "key_points": key_points,
        "action_items": action_items,
        "summary": summary,
        "participants": participants,
        "created_at": new_meeting.created_at,
        "language": target_language.lower()
    }


def getMeetings(db: Session):
    try:
        meetings = db.query(Meeting).all()
        return meetings
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch meetings: {str(e)}")


def getMeeting(id: int, db: Session, language: str = "english"):
    try:
        meeting = db.query(Meeting).filter(Meeting.id == id).first()
        if not meeting:
            raise HTTPException(status_code=404, detail="Meeting not found")
        
        transcript = None
        if meeting.transcript_id:
            transcription = db.query(Transcription).filter(Transcription.id == meeting.transcript_id).first()
            if transcription:
                transcript = transcription.transcript
        
        # Parse JSON fields
        key_points = json.loads(meeting.key_points) if meeting.key_points else []
        action_items = json.loads(meeting.action_items) if meeting.action_items else []
        participants = json.loads(meeting.participants) if meeting.participants else []
        
        meeting_response = {
            "id": meeting.id,
            "filename": meeting.filename if hasattr(meeting, 'filename') else None,
            "created_at": meeting.created_at,
            "title": meeting.title,
            "summary": meeting.summary,
            "duration": meeting.duration if hasattr(meeting, 'duration') else None,
            "participants": participants,
            "key_points": key_points,
            "action_items": action_items,
            "transcript": transcript,
            "language": "english"  # default language
        }
        
        # Translate if requested language is different
        if language.lower() != "english":
            meeting_response = translate_meeting_content(meeting_response, language)
        
        return meeting_response
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch meeting: {str(e)}")


def generate_meeting_report_pdf(meeting_id: int, db: Session, language: str = "english"):
    """Generate PDF report for a meeting with specified language"""
    try:
        meeting_data = getMeeting(meeting_id, db, language)
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_filename = temp_file.name
        temp_file.close()
        
        doc = SimpleDocTemplate(temp_filename, pagesize=A4, topMargin=1*inch, bottomMargin=1*inch)
        styles = getSampleStyleSheet()
        
        # Register Unicode font for Marathi support
        font_name = 'Helvetica'
        if language.lower() == "marathi":
            try:
                font_path = "fonts/NotoSansDevanagari-Regular.ttf"
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('DevanagariFont', font_path))
                    font_name = 'DevanagariFont'
                    print("Marathi font loaded successfully")
                else:
                    print(f"Warning: Font not found at {font_path}. Marathi text may not display correctly.")
            except Exception as font_error:
                print(f"Font registration error: {font_error}")
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue,
            fontName=font_name
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.darkblue,
            fontName=font_name
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            fontName=font_name
        )
        
        content = []
        
        # Language-specific labels
        if language.lower() == "marathi":
            report_title = "मीटिंग अहवाल"
            meeting_title_label = "मीटिंग शीर्षक:"
            date_label = "प्रक्रिया केलेली तारीख:"
            file_label = "फाइलचे नाव:"
            duration_label = "कालावधी:"
            summary_label = "कार्यकारी सारांश"
            participants_label = "सहभागी"
            key_points_label = "मुख्य मुद्दे"
            action_items_label = "कृती आयटम"
            transcript_label = "संपूर्ण ट्रान्सक्रिप्ट"
        else:
            report_title = "Meeting Report"
            meeting_title_label = "Meeting Title:"
            date_label = "Date Processed:"
            file_label = "File Name:"
            duration_label = "Duration:"
            summary_label = "Executive Summary"
            participants_label = "Participants"
            key_points_label = "Key Points"
            action_items_label = "Action Items"
            transcript_label = "Full Transcript"
        
        content.append(Paragraph(report_title, title_style))
        content.append(Spacer(1, 20))
        
        meeting_info_data = [
            [meeting_title_label, str(meeting_data.get("title", "N/A"))],
            [date_label, meeting_data.get("created_at", "").strftime("%Y-%m-%d %H:%M:%S") if meeting_data.get("created_at") else "N/A"],
            [file_label, str(meeting_data.get("filename", "N/A"))],
            [duration_label, str(meeting_data.get("duration", "N/A"))]
        ]
        
        meeting_info_table = Table(meeting_info_data, colWidths=[2*inch, 4*inch])
        meeting_info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        content.append(meeting_info_table)
        content.append(Spacer(1, 20))
        
        if meeting_data.get("summary"):
            content.append(Paragraph(summary_label, heading_style))
            summary_text = str(meeting_data["summary"])
            content.append(Paragraph(summary_text, normal_style))
            content.append(Spacer(1, 20))
        
        participants = meeting_data.get("participants", [])
        if participants:
            content.append(Paragraph(participants_label, heading_style))
            participants_text = ", ".join([str(p) for p in participants])
            content.append(Paragraph(participants_text, normal_style))
            content.append(Spacer(1, 20))
        
        key_points = meeting_data.get("key_points", [])
        if key_points:
            content.append(Paragraph(key_points_label, heading_style))
            for point in key_points:
                point_text = f"• {str(point)}"
                try:
                    content.append(Paragraph(point_text, normal_style))
                except Exception as e:
                    print(f"Error adding key point: {e}")
            content.append(Spacer(1, 20))
        
        action_items = meeting_data.get("action_items", [])
        if action_items:
            content.append(Paragraph(action_items_label, heading_style))
            for item in action_items:
                item_text = f"• {str(item)}"
                try:
                    content.append(Paragraph(item_text, normal_style))
                except Exception as e:
                    print(f"Error adding action item: {e}")
            content.append(Spacer(1, 20))
        
        if meeting_data.get("transcript"):
            content.append(PageBreak())
            content.append(Paragraph(transcript_label, heading_style))
            
            transcript = str(meeting_data["transcript"])
            sentences = transcript.split('. ')
            paragraphs = []
            
            for i in range(0, len(sentences), 3):
                paragraph = '. '.join(sentences[i:i+3])
                if paragraph and not paragraph.endswith('.'):
                    paragraph += '.'
                paragraphs.append(paragraph)
            
            for paragraph in paragraphs:
                if paragraph.strip():
                    try:
                        content.append(Paragraph(paragraph, normal_style))
                        content.append(Spacer(1, 10))
                    except Exception as para_error:
                        print(f"Error adding paragraph: {para_error}")
                        continue
        
        doc.build(content)
        print(f"PDF generated successfully: {temp_filename}")
        
        return temp_filename
        
    except Exception as e:
        print(f"PDF Generation Error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF report: {str(e)}")


def generate_meeting_report_docx(meeting_id: int, db: Session, language: str = "english"):
    """Generate DOCX report for a meeting with specified language"""
    try:
        meeting_data = getMeeting(meeting_id, db, language)
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_filename = temp_file.name
        temp_file.close()
        
        doc = Document()
        
        # Language-specific labels
        if language.lower() == "marathi":
            report_title = "मीटिंग अहवाल"
            meeting_info_label = "मीटिंग माहिती"
            meeting_title_label = "मीटिंग शीर्षक:"
            date_label = "प्रक्रिया केलेली तारीख:"
            file_label = "फाइलचे नाव:"
            duration_label = "कालावधी:"
            summary_label = "कार्यकारी सारांश"
            participants_label = "सहभागी"
            key_points_label = "मुख्य मुद्दे"
            action_items_label = "कृती आयटम"
            transcript_label = "संपूर्ण ट्रान्सक्रिप्ट"
        else:
            report_title = "Meeting Report"
            meeting_info_label = "Meeting Information"
            meeting_title_label = "Meeting Title:"
            date_label = "Date Processed:"
            file_label = "File Name:"
            duration_label = "Duration:"
            summary_label = "Executive Summary"
            participants_label = "Participants"
            key_points_label = "Key Points"
            action_items_label = "Action Items"
            transcript_label = "Full Transcript"
        
        title = doc.add_heading(report_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading(meeting_info_label, level=1)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = table.rows[0].cells
        cells[0].text = meeting_title_label
        cells[1].text = str(meeting_data.get("title", "N/A"))
        
        cells = table.rows[1].cells
        cells[0].text = date_label
        created_at = meeting_data.get("created_at")
        if created_at:
            try:
                cells[1].text = created_at.strftime("%Y-%m-%d %H:%M:%S")
            except:
                cells[1].text = str(created_at)
        else:
            cells[1].text = "N/A"
        
        cells = table.rows[2].cells
        cells[0].text = file_label
        cells[1].text = str(meeting_data.get("filename", "N/A"))
        
        cells = table.rows[3].cells
        cells[0].text = duration_label
        cells[1].text = str(meeting_data.get("duration", "N/A"))
        
        if meeting_data.get("summary"):
            doc.add_heading(summary_label, level=1)
            doc.add_paragraph(str(meeting_data["summary"]))
        
        participants = meeting_data.get("participants", [])
        if participants:
            doc.add_heading(participants_label, level=1)
            participants_text = ", ".join([str(p) for p in participants])
            doc.add_paragraph(participants_text)
        
        key_points = meeting_data.get("key_points", [])
        if key_points:
            doc.add_heading(key_points_label, level=1)
            for point in key_points:
                p = doc.add_paragraph()
                p.add_run(f"• {str(point)}")
        
        action_items = meeting_data.get("action_items", [])
        if action_items:
            doc.add_heading(action_items_label, level=1)
            for item in action_items:
                p = doc.add_paragraph()
                p.add_run(f"• {str(item)}")
        
        if meeting_data.get("transcript"):
            doc.add_page_break()
            doc.add_heading(transcript_label, level=1)
            
            transcript = str(meeting_data["transcript"])
            sentences = transcript.split('. ')
            paragraphs = []
            
            for i in range(0, len(sentences), 3):
                paragraph = '. '.join(sentences[i:i+3])
                if paragraph and not paragraph.endswith('.'):
                    paragraph += '.'
                paragraphs.append(paragraph)
            
            for paragraph in paragraphs:
                if paragraph.strip():
                    try:
                        doc.add_paragraph(paragraph)
                    except Exception as para_error:
                        print(f"Error adding paragraph: {para_error}")
                        continue
        
        doc.save(temp_filename)
        print(f"DOCX generated successfully: {temp_filename}")
        
        return temp_filename
        
    except Exception as e:
        print(f"DOCX Generation Error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX report: {str(e)}")


def download_meeting_report(meeting_id: int, format_type: str, db: Session, language: str = "english"):
    """Download meeting report in PDF or DOCX format with specified language"""
    try:
        print(f"Generating {format_type} report for meeting {meeting_id} in {language}")
        
        meeting_data = getMeeting(meeting_id, db, language)
        meeting_title = meeting_data.get("title", "Meeting Report").replace(" ", "_")
        
        if format_type.lower() == "pdf":
            file_path = generate_meeting_report_pdf(meeting_id, db, language)
            filename = f"{meeting_title}_Report.pdf"
            media_type = "application/pdf"
        elif format_type.lower() == "docx":
            file_path = generate_meeting_report_docx(meeting_id, db, language)
            filename = f"{meeting_title}_Report.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            raise HTTPException(status_code=400, detail="Invalid format type. Use 'pdf' or 'docx'")
        
        # Return FileResponse without background cleanup
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type=media_type
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Download report error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to download report: {str(e)}")